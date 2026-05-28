from docx import Document
from docx.shared import Pt
import os
from docx.shared import Inches

CONTENT_TITLE = "BAB IV IMPLEMENTASI DAN PENGUJIAN"

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for r in rows:
        row_cells = table.add_row().cells
        for i, c in enumerate(r):
            row_cells[i].text = str(c)
    return table

def build_document(output_path):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    add_heading(doc, CONTENT_TITLE, level=1)

    add_heading(doc, 'IV.1 Implementasi Sistem', level=2)
    add_paragraph(doc, 'Implementasi sistem merupakan tahap penerapan hasil perancangan ke dalam bentuk aplikasi berbasis web. Sistem COCONEXUS dikembangkan menggunakan arsitektur 3-lapis dengan backend Node.js (Express) dan database MySQL serta frontend berbasis Vue 3. Sistem ini dirancang untuk mendukung repository pembelajaran pemanfaatan limbah kelapa dengan fitur publikasi artikel, manajemen kategori, komentar, unggah media, statistik artikel, dan pengelolaan pengguna.')

    add_paragraph(doc, 'Sistem memiliki tiga jenis aktor utama: pengguna umum, moderator (content moderator), dan admin. Pengguna umum dapat mengakses fitur publik tanpa harus login, seperti daftar artikel, detail artikel, komentari, dan halaman informasi. Moderator memiliki hak akses untuk membuat, menyunting, dan mempublikasikan artikel dengan workflow (draft → revision → published), mengelola kategori pembelajaran, serta meninjau komentar. Admin memiliki hak akses penuh untuk mengelola data pengguna, kategori, artikel, komentar, dan melihat log audit serta statistik.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Tabel IV.1.1 Tabel implementasi sistem')
    headers = ['No', 'Komponen', 'Hasil Implementasi']
    rows = [
        [1, 'Framework', 'Backend: Node.js + Express; Frontend: Vue 3'],
        [2, 'Basis Data', 'MySQL dengan ORM Sequelize'],
        [3, 'Akses Sistem', 'Akses melalui browser; peran: publik, moderator, admin'],
        [4, 'Arsitektur', 'Pola Model-View-Controller (MVC) dan REST API'],
        [5, 'Hak Akses', 'Pembatasan fitur berdasarkan role (JWT + middleware)'],
    ]
    add_table(doc, headers, rows)
    # tambahan: detail kemampuan pengelolaan artikel
    doc.add_paragraph('Sistem dapat membuat, menyunting, dan mempublikasikan artikel dengan workflow (draft → revision → published).')
    doc.add_paragraph('Sistem dapat mengelola kategori pembelajaran secara otomatis saat pembuatan atau pembaruan artikel.', style='List Bullet')
    doc.add_paragraph('Sistem dapat menangani detail artikel dengan media pendukung seperti gambar, video, dan dokumen.', style='List Bullet')
    doc.add_paragraph('Sistem dapat melacak jumlah views artikel untuk mengukur popularitas konten.', style='List Bullet')

    add_heading(doc, 'IV.2 Implementasi Database', level=2)
    add_paragraph(doc, 'Implementasi database dilakukan menggunakan MySQL. Struktur database dibuat berdasarkan rancangan ERD yang telah disusun pada tahap perancangan. Database menyimpan data pengguna, profil, role, kategori, artikel, detail artikel, media, komentar, views, product cards, dan audit logs. Relasi antar tabel dibangun menggunakan primary key dan foreign key agar data tersimpan dan terhubung secara terstruktur.')

    add_paragraph(doc, 'Tabel IV.2.1 Tabel implementasi database')
    headers = ['No', 'Nama Tabel', 'Fungsi Implementasi']
    rows = [
        [1, 'roles', 'Menyimpan data peran atau hak akses pengguna'],
        [2, 'users', 'Menyimpan data akun pengguna sistem'],
        [3, 'user_profiles', 'Menyimpan profil pengguna (bio, avatar)'],
        [4, 'category_tags', 'Menyimpan kategori pembelajaran'],
        [5, 'articles', 'Menyimpan data utama artikel'],
        [6, 'article_details', 'Menyimpan konten detail artikel (sections, sources)'],
        [7, 'article_media', 'Menyimpan media pendukung artikel (gambar, video, dokumen)'],
        [8, 'comments', 'Menyimpan komentar bersarang pada artikel'],
        [9, 'article_views', 'Menyimpan hitungan views per artikel'],
        [10, 'product_cards', 'Menyimpan data product card terkait artikel'],
        [11, 'audit_logs', 'Menyimpan riwayat aktivitas pengguna dalam sistem'],
    ]
    add_table(doc, headers, rows)

    add_heading(doc, 'IV.3 Implementasi Antarmuka', level=2)
    add_paragraph(doc, 'Implementasi antarmuka dilakukan berdasarkan rancangan tampilan. Antarmuka sistem berbasis web dan disesuaikan dengan hak akses pengguna sehingga pengguna umum, moderator, dan admin memperoleh menu berbeda.')

    add_paragraph(doc, 'Gambar referensi (ditunjukkan di laporan sebagai bukti implementasi):')
    add_paragraph(doc, '- Halaman Login')
    add_paragraph(doc, '- Dashboard user')
    add_paragraph(doc, '- Form pembuatan artikel')
    add_paragraph(doc, '- Form upload media')
    add_paragraph(doc, '- Halaman detail artikel dengan komentar')
    add_paragraph(doc, '- Dashboard admin (manajemen artikel, user, kategori)')

    add_paragraph(doc, 'Tabel IV.3.1 Tabel implementasi antarmuka')
    headers = ['No', 'Halaman', 'Aktor', 'Keterangan']
    rows = [
        [1, 'Login', 'Moderator dan Admin', 'Digunakan untuk masuk ke dashboard internal sesuai role'],
        [2, 'Landing Page / Artikel Publik', 'Pengguna Umum', 'Menampilkan daftar artikel, kategori, dan pencarian'],
        [3, 'Form Pembuatan Artikel', 'Moderator dan Admin', 'Form untuk membuat artikel, mengatur kategori, upload media'],
        [4, 'Detail Artikel', 'Semua pengguna', 'Menampilkan konten, media, dan komentar bersarang'],
        [5, 'Dashboard Admin', 'Admin', 'Manajemen data user, artikel, kategori, dan audit logs'],
    ]
    add_table(doc, headers, rows)

    add_heading(doc, 'IV.4 Implementasi Fitur', level=2)
    add_paragraph(doc, 'Implementasi fitur disesuaikan kebutuhan sistem:')

    add_paragraph(doc, 'Tabel IV.4.1 Tabel implementasi fitur')
    headers = ['No', 'Fitur', 'Hasil Implementasi']
    rows = [
        [1, 'Manajemen Artikel', 'Moderator/Admin dapat membuat, edit, dan publish artikel (draft → revision → published)'],
        [2, 'Manajemen Kategori', 'Otomatis membuat/memperbarui kategori saat penyimpanan artikel'],
        [3, 'Media & Upload', 'Menangani upload gambar, video, dokumen untuk artikel dan avatar pengguna'],
        [4, 'Komentar Bersarang', 'Mendukung komentar bersarang dengan parent_id'],
        [5, 'Statistik & Views', 'Melacak jumlah views artikel untuk metrik popularitas'],
        [6, 'Manajemen Pengguna', 'Admin dapat menambah, ubah, soft delete pengguna'],
        [7, 'Audit Log', 'Mencatat aktivitas penting di tabel audit_logs'],
    ]
    add_table(doc, headers, rows)

    add_heading(doc, 'IV.5 Pengujian Sistem', level=2)
    add_paragraph(doc, 'Pengujian dilakukan untuk memastikan fitur berjalan sesuai kebutuhan. Metode pengujian meliputi black box testing (Newman untuk API), pengujian integrasi backend, dan smoke test pada UI.')

    add_heading(doc, 'IV.5.1 Pengujian Black Box', level=3)
    add_paragraph(doc, 'Pengujian black box dilakukan dengan koleksi Postman dan dieksekusi menggunakan Newman. Pengujian mencakup endpoint publik (read articles), endpoint internal (create/update/delete article), upload media, komentar, autentikasi, dan otorisasi role.')

    add_heading(doc, 'IV.5.2 Pengujian User Acceptance Test / UAT', level=3)
    add_paragraph(doc, 'UAT disusun berdasarkan peran: Pengguna umum, Moderator, Admin. Skenario UAT meliputi pendaftaran, login, pembuatan artikel, publikasi, penambahan komentar, dan pengelolaan data oleh admin.')

    add_heading(doc, 'IV.6 Hasil Pengujian', level=2)
    add_paragraph(doc, 'Hasil pengujian backend menunjukkan seluruh skenario API yang diuji berhasil. Newman report dan ringkasan pengujian disimpan dalam folder `docs/newman-results`.')

    add_paragraph(doc, '- Black Box menggunakan Newman: 23 request dan 36 assertion — semua berhasil.')
    add_paragraph(doc, '- UI Smoke Test: 18 skenario — semua berhasil.')

    add_heading(doc, 'IV.7 Pembahasan', level=2)
    add_paragraph(doc, 'Berdasarkan implementasi dan pengujian, sistem COCONEXUS telah memenuhi kebutuhan fungsional utama. Sistem memungkinkan pembuatan dan publikasi artikel, manajemen kategori, upload media, komentar bersarang, dan pelacakan view. Pengujian otomatis dan manual menunjukkan fitur berjalan sesuai skenario yang diuji.')

    add_paragraph(doc, '\nLampiran: Lampiran bukti implementasi (screenshot halaman, newman reports, dan file migrasi) disimpan di folder `docs/` dan `backend/migrations`.')

    # Lampiran: embed screenshot and list newman report files if available
    add_heading(doc, 'Lampiran - Bukti Pengujian', level=2)
    newman_dir = os.path.join(repo_root, 'docs', 'newman-results')
    screenshot_path = os.path.join(newman_dir, 'ringkasan-hasil-pengujian.png')
    if os.path.exists(screenshot_path):
        try:
            doc.add_paragraph('Gambar IV.6.1 Ringkasan Hasil Pengujian (screenshot):')
            doc.add_picture(screenshot_path, width=Inches(6))
        except Exception:
            doc.add_paragraph('Gambar ringkasan tidak dapat disematkan, lihat file di docs/newman-results/')
    else:
        doc.add_paragraph('Screenshot pengujian tidak ditemukan di docs/newman-results/')

    # list available report files
    doc.add_paragraph('Daftar file laporan Newman dan dokumentasi pengujian:')
    if os.path.exists(newman_dir):
        for fname in sorted(os.listdir(newman_dir)):
            fpath = os.path.join('docs', 'newman-results', fname)
            doc.add_paragraph(f'- {fname}: lihat {fpath}')
    else:
        doc.add_paragraph('Folder docs/newman-results tidak ditemukan.')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print('Saved:', output_path)

if __name__ == '__main__':
    repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    out = os.path.join(repo_root, 'docs', 'BAB-IV-COCONEXUS-implementasi_v2.docx')
    build_document(out)
