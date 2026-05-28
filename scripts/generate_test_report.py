from docx import Document
from docx.shared import Pt, Inches
import os

repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
out = os.path.join(repo_root, 'docs', 'BAB-IV-TEST-REPORT.docx')
os.makedirs(os.path.dirname(out), exist_ok=True)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

doc.add_heading('Laporan Hasil Pengujian COCONEXUS', level=1)

doc.add_heading('1. Ringkasan', level=2)
doc.add_paragraph('Tes yang dijalankan: backend unit/integration tests, Newman API reports (existing), backend smoke test, frontend build (UI smoke).')

# Backend tests
backend_test_file = os.path.join(repo_root, 'docs', 'test', 'backend-tests.txt')
if os.path.exists(backend_test_file):
    doc.add_heading('2. Hasil Tes Backend', level=2)
    with open(backend_test_file, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    # sanitize control characters
    def sanitize(s):
        return ''.join(ch for ch in s if (ch == '\n' or ch == '\t' or ord(ch) >= 32))
    content = sanitize(content)
    # add as paragraph blocks
    for chunk in content.splitlines():
        doc.add_paragraph(chunk)
else:
    doc.add_paragraph('Hasil tes backend tidak ditemukan.')

# Smoke test
smoke_file = os.path.join(repo_root, 'docs', 'test', 'smoke-test.txt')
if os.path.exists(smoke_file):
    doc.add_heading('3. Hasil Smoke Test (Backend)', level=2)
    with open(smoke_file, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    content = ''.join(ch for ch in content if (ch == '\n' or ch == '\t' or ord(ch) >= 32))
    for chunk in content.splitlines():
        doc.add_paragraph(chunk)
else:
    doc.add_paragraph('Hasil smoke test tidak ditemukan.')

# Frontend build log
frontend_file = os.path.join(repo_root, 'docs', 'test', 'frontend-build.txt')
if os.path.exists(frontend_file):
    doc.add_heading('4. Hasil Build Frontend (UI Smoke)', level=2)
    with open(frontend_file, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    content = ''.join(ch for ch in content if (ch == '\n' or ch == '\t' or ord(ch) >= 32))
    for chunk in content.splitlines():
        doc.add_paragraph(chunk)
else:
    doc.add_paragraph('Hasil build frontend tidak ditemukan.')

# Newman reports listing and embed screenshot
newman_dir = os.path.join(repo_root, 'docs', 'newman-results')
doc.add_heading('5. Newman Reports', level=2)
if os.path.exists(newman_dir):
    for fname in sorted(os.listdir(newman_dir)):
        doc.add_paragraph(f'- {fname}')
    # try to embed screenshot
    screenshot = os.path.join(newman_dir, 'ringkasan-hasil-pengujian.png')
    if os.path.exists(screenshot):
        doc.add_heading('6. Screenshot Ringkasan Pengujian', level=2)
        try:
            doc.add_picture(screenshot, width=Inches(6))
        except Exception as e:
            doc.add_paragraph('Gagal menyematkan screenshot: ' + str(e))
else:
    doc.add_paragraph('Folder newman-results tidak ditemukan.')

# save
doc.save(out)
print('Saved:', out)
